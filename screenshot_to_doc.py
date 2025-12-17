import os
import sys
import tempfile
import time
from docx import Document
from docx.shared import Inches


def main():
    print("Welcome to the Screenshot to Doc Tool")
    print("-------------------------------------")

    # 1. Ask for document name
    doc_name = input("Enter output document name (default: screenshots.docx): ").strip()
    if not doc_name:
        doc_name = "screenshots.docx"

    if not doc_name.endswith(".docx"):
        doc_name += ".docx"

    # 2. Open or Create Document
    if os.path.exists(doc_name):
        print(f"Opening existing document: {doc_name}")
        try:
            doc = Document(doc_name)
        except Exception as e:
            print(f"Error opening document: {e}")
            sys.exit(1)
    else:
        print(f"Creating new document: {doc_name}")
        doc = Document()

    print("\nControls:")
    print("  [Enter]  : Capture Screenshot")
    print("  [Ctrl+C] : Save and Exit")
    print("-------------------------------------")

    try:
        while True:
            # Wait for user input to capture
            input(
                f"\nPress Enter to capture screenshot #{len(doc.inline_shapes) + 1}...")

            # Use a temporary file for the screenshot
            # delete=False is needed because we close it before gnome-screenshot writes to it (os.system)
            # Actually simplest way is just generation a temp path string
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
                tmp_img_path = tf.name

            # Capture screenshot
            # gnome-screenshot -f <filename> captures the entire screen
            cmd = f"gnome-screenshot -f '{tmp_img_path}'"
            ret_code = os.system(cmd)

            if ret_code != 0:
                print("Error: Failed to capture screenshot. Is gnome-screenshot installed?")
                # Clean up if file was created empty or not created
                if os.path.exists(tmp_img_path):
                    os.remove(tmp_img_path)
                continue

            # Verify file exists and has size
            if not os.path.exists(tmp_img_path) or os.path.getsize(tmp_img_path) == 0:
                print("Error: Screenshot file is empty or missing.")
                continue

            try:
                # Add to doc
                # Add a page break if it's not the first one? Or just append.
                # User usually wants one by one. I'll just append.
                if len(doc.inline_shapes) > 0:
                    doc.add_paragraph(
                        f"\nScreenshot captured at {time.strftime('%Y-%m-%d %H:%M:%S')}")
                else:
                    doc.add_paragraph(
                        f"Screenshot captured at {time.strftime('%Y-%m-%d %H:%M:%S')}")

                doc.add_picture(tmp_img_path, width=Inches(6.5))
                print(f"Screenshot captured and added!")
            except Exception as e:
                print(f"Error adding image to document: {e}")
            finally:
                # Cleanup temp file
                if os.path.exists(tmp_img_path):
                    os.remove(tmp_img_path)

    except KeyboardInterrupt:
        print("\n\nStopping...")

    # Save document
    print(f"Saving {doc_name}...")
    try:
        doc.save(doc_name)
        print("Done!")
    except Exception as e:
        print(f"Error saving document: {e}")
        # Try saving with timestamp if locked
        timestamp_name = f"screenshots_{int(time.time())}.docx"
        print(f"Trying to save as {timestamp_name}...")
        doc.save(timestamp_name)
        print("Done!")


if __name__ == "__main__":
    main()
